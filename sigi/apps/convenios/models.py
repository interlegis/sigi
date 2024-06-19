import re
import requests
from difflib import SequenceMatcher
from hashlib import md5
from pathlib import Path
from django.db import models
from django.db.models import Q, F
from django.contrib.sites.shortcuts import get_current_site
from django.core.exceptions import ValidationError, NON_FIELD_ERRORS
from django.core.mail import send_mail
from django.core.validators import FileExtensionValidator
from django.template import Template, Context
from django.template.exceptions import TemplateSyntaxError
from django.urls import reverse
from django.utils import timezone
from django.utils.formats import date_format
from django.utils.translation import gettext as _
from django_weasyprint.utils import django_url_fetcher
from docx import Document
from tinymce.models import HTMLField
from weasyprint import HTML
from sigi.apps.contatos.models import Municipio, UnidadeFederativa
from sigi.apps.parlamentares.models import Parlamentar
from sigi.apps.utils import to_ascii
from sigi.apps.casas.models import Funcionario, Orgao
from sigi.apps.servidores.models import Servidor, Servico
from sigi.apps.utils import editor_help, mask_cnpj


class Projeto(models.Model):
    OFICIO_HELP = editor_help(
        "texto_oficio",
        [
            ("casa", Orgao),
            ("presidente", Parlamentar),
            ("contato", Funcionario),
            ("casa.municipio", Municipio),
            ("casa.municipio.uf", UnidadeFederativa),
            ("data", _("Data atual")),
            ("doravante", _("CÂMARA ou ASSEMBLEIA")),
        ],
    )
    MINUTA_HELP = editor_help(
        "modelo_minuta",
        [
            ("casa", Orgao),
            ("presidente", Parlamentar),
            ("contato", Funcionario),
            ("casa.municipio", Municipio),
            ("casa.municipio.uf", UnidadeFederativa),
            ("data", _("Data atual")),
            ("ente", _("Ente da federação (município/estado)")),
            ("doravante", _("CÂMARA ou ASSEMBLEIA")),
        ],
    )
    nome = models.CharField(max_length=50)
    sigla = models.CharField(max_length=10)
    texto_oficio = HTMLField(
        _("texto do ofício"), blank=True, help_text=OFICIO_HELP
    )
    modelo_minuta = models.FileField(
        _("Modelo de minuta"),
        blank=True,
        help_text=MINUTA_HELP,
        upload_to="convenios/minutas/",
        validators=[
            FileExtensionValidator(
                [
                    "docx",
                ]
            ),
        ],
    )

    def __str__(self):
        return self.sigla

    class Meta:
        ordering = ("nome",)

    def gerar_oficio(self, file_object, casa, presidente, contato, path):
        texto = self.texto_oficio
        template_string = (
            '{% extends "convenios/oficio_padrao.html" %}'
            "{% load pdf %}"
            f"{{% block text_body %}}{texto}{{% endblock %}}"
        )
        context = Context(
            {
                "casa": casa,
                "presidente": presidente,
                "contato": contato,
                "data": timezone.localdate(),
                "doravante": casa.tipo.nome.split(" ")[0],
            }
        )
        string = Template(template_string).render(context)
        pdf = HTML(
            string=string,
            url_fetcher=django_url_fetcher,
            encoding="utf-8",
            base_url=path,
        )
        file_name = Path(file_object.path)
        if not file_name.parent.exists():
            file_name.parent.mkdir(parents=True, exist_ok=True)
        if file_object.closed:
            file_object.open(mode="wb")
        pdf.write_pdf(target=file_object)
        file_object.flush()

    def gerar_minuta(self, file_path, casa, presidente, contato):
        doc = Document(self.modelo_minuta.path)

        if casa.tipo.sigla == "CM":
            ente = (
                f"Município de {casa.municipio.nome}, "
                f"{casa.municipio.uf.sigla}"
            )
        else:
            ente = f"Estado de {casa.municipio.uf.nome}"

        doc_context = Context(
            {
                "casa": casa,
                "presidente": presidente,
                "contato": contato,
                "data": timezone.localdate(),
                "ente": ente,
                "doravante": casa.tipo.nome.split(" ")[0],
            }
        )

        self.processa_paragrafos(doc.paragraphs, doc_context)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.processa_paragrafos(
                        cell.paragraphs,
                        doc_context,
                    )
        file_path = Path(file_path)
        if not file_path.parent.exists():
            file_path.mkdir(parents=True, exist_ok=True)
        doc.save(file_path)

    def processa_paragrafos(self, paragrafos, context):
        for paragrafo in paragrafos:
            run_final = None
            for run in paragrafo.runs:
                if run_final is None:
                    run_final = run
                else:
                    run_final.text += run.text
                    run.text = ""
                if run_final.text.count("{{") != run_final.text.count("}}"):
                    continue
                try:
                    run_final.text = Template(run_final.text).render(context)
                    run_final = None
                except TemplateSyntaxError:
                    pass


class StatusConvenio(models.Model):
    nome = models.CharField(max_length=100)
    cancela = models.BooleanField(_("Cancela o convênio"), default=False)

    class Meta:
        ordering = ("nome",)
        verbose_name = _("Estado de convenios")
        verbose_name_plural = _("Estados de convenios")

    def __str__(self):
        return self.nome


class TipoSolicitacao(models.Model):
    nome = models.CharField(max_length=100)

    class Meta:
        ordering = ("nome",)
        verbose_name = _("tipo de solicitação")
        verbose_name_plural = _("Tipos de solicitação")

    def __str__(self):
        return self.nome


class Convenio(models.Model):
    casa_legislativa = models.ForeignKey(
        "casas.Orgao",
        on_delete=models.PROTECT,
        verbose_name=_("órgão conveniado"),
    )
    projeto = models.ForeignKey(
        Projeto, on_delete=models.PROTECT, verbose_name=_("Tipo de Convenio")
    )
    # numero designado pelo Senado Federal para o convênio
    num_processo_sf = models.CharField(
        _("número do processo SF (Senado Federal)"),
        max_length=20,
        blank=True,
        help_text=_(
            "Formatos:<br/>Antigo: <em>XXXXXX/XX-X</em>.<br/><em>SIGAD: XXXXX.XXXXXX/XXXX-XX</em>"
        ),
    )
    # link_processo_stf = ('get_sigad_url')
    num_convenio = models.CharField(
        _("número do convênio"), max_length=10, blank=True
    )
    id_contrato_gescon = models.CharField(
        _("ID do contrato no Gescon"),
        max_length=20,
        blank=True,
        default="",
        editable=False,
    )
    data_sigi = models.DateField(
        _("data de cadastro no SIGI"), blank=True, null=True, auto_now_add=True
    )
    data_sigad = models.DateField(
        _("data de cadastro no SIGAD"), null=True, blank=True
    )
    data_solicitacao = models.DateField(
        _("data do e-mail de solicitação"), null=True, blank=True
    )
    tipo_solicitacao = models.ForeignKey(
        TipoSolicitacao,
        on_delete=models.PROTECT,
        null=True,
        blank=True,
        verbose_name=_("tipo de solicitação"),
    )
    status = models.ForeignKey(
        StatusConvenio,
        on_delete=models.SET_NULL,
        verbose_name=_("estado atual"),
        null=True,
        blank=True,
    )
    acompanha = models.ForeignKey(
        Servidor,
        on_delete=models.SET_NULL,
        related_name="convenios_acompanhados",
        verbose_name=_("acompanhado por"),
        null=True,
        blank=True,
    )
    observacao = models.TextField(
        _("observações"),
        null=True,
        blank=True,
    )
    servico_gestao = models.ForeignKey(
        Servico,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name="convenios_geridos",
        verbose_name=_("serviço de gestão"),
    )
    servidor_gestao = models.ForeignKey(
        Servidor,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        verbose_name=_("servidor de gestão"),
    )
    data_adesao = models.DateField(
        _("aderidas"),
        null=True,
        blank=True,
    )
    data_retorno_assinatura = models.DateField(
        _("data início vigência"),
        null=True,
        blank=True,
        help_text=_("Convênio firmado."),
    )
    data_termino_vigencia = models.DateField(
        _("Data término vigência"),
        null=True,
        blank=True,
        help_text=_("Término da vigência do convênio."),
    )
    data_pub_diario = models.DateField(
        _("data da publicação no Diário Oficial"), null=True, blank=True
    )
    data_termo_aceite = models.DateField(
        _("equipadas"),
        null=True,
        blank=True,
        help_text=_("Equipamentos recebidos."),
    )
    data_devolucao_via = models.DateField(
        _("data de devolução da via"),
        null=True,
        blank=True,
        help_text=_(
            "Data de devolução da via do convênio à Câmara Municipal."
        ),
    )
    data_postagem_correio = models.DateField(
        _("data postagem correio"),
        null=True,
        blank=True,
    )
    data_devolucao_sem_assinatura = models.DateField(
        _("data de devolução por falta de assinatura"),
        null=True,
        blank=True,
        help_text=_("Data de devolução por falta de assinatura"),
    )
    data_retorno_sem_assinatura = models.DateField(
        _("data do retorno sem assinatura"),
        null=True,
        blank=True,
        help_text=_("Data do retorno do convênio sem assinatura"),
    )
    data_extincao = models.DateField(
        _("data de extinção/desistência"), null=True, blank=True
    )
    motivo_extincao = models.TextField(
        _("motivo da extinção/desistência"), blank=True
    )
    conveniada = models.BooleanField(default=False)
    equipada = models.BooleanField(default=False)
    atualizacao_gescon = models.DateTimeField(
        _("Data de atualização pelo Gescon"), blank=True, null=True
    )
    erro_gescon = models.BooleanField(
        _("erro no Gescon"),
        max_length=1,
        default=False,
    )
    observacao_gescon = models.TextField(
        _("Observações da atualização do Gescon"), blank=True
    )

    def get_status(self):
        if self.status and self.status.cancela:
            return _("Cancelado")

        if self.data_extincao:
            return _("Extinto")

        if self.data_retorno_assinatura is not None:
            if self.data_termino_vigencia is not None:
                if timezone.localdate() >= self.data_termino_vigencia:
                    return _("Vencido")
            return _("Vigente")

        if (
            self.data_retorno_assinatura is None
            and self.data_devolucao_sem_assinatura is None
            and self.data_retorno_sem_assinatura is None
        ):
            return _("Pendente")
        if (
            self.data_devolucao_sem_assinatura is not None
            or self.data_retorno_sem_assinatura is not None
        ):
            return _("Desistência")

        return _("Indefinido")

    def link_sigad(self, obj):
        if obj.pk is None:
            return ""
        return obj.get_sigad_url()

    def get_sigad_url(self, display_type="numero"):
        m = re.match(
            r"(?P<orgao>00100|00200)\.(?P<sequencial>\d{6})/(?P<ano>\d{4})-\d{2}",
            self.num_processo_sf,
        )
        if m:
            orgao, sequencial, ano = m.groups()
            if display_type == "numero":
                display = self.num_processo_sf
            else:
                display = "<i class='material-icons'>visibility</i>"
            return (
                f'<a href="https://intra.senado.leg.br/sigad/novo/protocolo/'
                f"impressao.asp?area=processo&txt_numero_orgao={orgao}"
                f'&txt_numero_sequencial={sequencial}&txt_numero_ano={ano}" '
                f'target="_blank">{display}</a>'
            )
        if display_type == "numero":
            return self.num_processo_sf
        else:
            return "<i class='material-icons'>visibility_off</i>"

    def get_url_gescon(self):
        if not self.id_contrato_gescon:
            return ""
        return (
            "https://adm.senado.gov.br/gestao-contratos/api/contratos"
            f"/buscaTexto/{self.id_contrato_gescon}"
        )

    def get_url_minuta(self):
        if self.id_contrato_gescon:
            return self.get_link_gescon()
        if self.anexo_set.exists():
            return self.anexo_set.first().arquivo.url
        return ""

    def clean(self):
        # Gertiq #184827
        if self.num_convenio:
            if (
                self.data_retorno_assinatura is None
                or self.data_termino_vigencia is None
            ):
                errors = {
                    NON_FIELD_ERRORS: ValidationError(
                        _(
                            "Um convênio vigente precisa ter as datas de "
                            "início e término de vigência"
                        )
                    )
                }
                if self.data_retorno_assinatura is None:
                    errors["data_retorno_assinatura"] = ValidationError(
                        _("Obrigatório para convênios vigentes")
                    )
                if self.data_termino_vigencia is None:
                    errors["data_termino_vigencia"] = ValidationError(
                        _("Obrigatório para convênios vigentes")
                    )
                raise ValidationError(errors)
        else:
            if (
                self.data_retorno_assinatura is not None
                or self.data_termino_vigencia is not None
            ):
                errors = {
                    NON_FIELD_ERRORS: ValidationError(
                        _(
                            "Um convênio pendente não pode ter datas de "
                            "início e término de vigência"
                        )
                    )
                }
                if self.data_retorno_assinatura is not None:
                    errors["data_retorno_assinatura"] = ValidationError(
                        _("Não pode ser preenchido para convênios pendentes")
                    )
                if self.data_termino_vigencia is not None:
                    errors["data_termino_vigencia"] = ValidationError(
                        _("Não pode ser preenchido para convênios pendentes")
                    )
                raise ValidationError(errors)

        return super().clean()

    def save(self, *args, **kwargs):
        self.conveniada = self.data_retorno_assinatura is not None
        self.equipada = self.data_termo_aceite is not None
        super().save(*args, **kwargs)

    class Meta:
        get_latest_by = "id"
        ordering = ("id",)
        verbose_name = _("convênio")

    def __str__(self):
        from django.conf import settings

        SDF = settings.SHORT_DATE_FORMAT
        number = self.num_convenio
        project = self.projeto.sigla
        if self.data_extincao:
            date = date_format(self.data_extincao, SDF)
            return _(f"{project} nº {number} extinto em {date}")
        if (self.data_retorno_assinatura is None) and (
            self.equipada and self.data_termo_aceite is not None
        ):
            date = date_format(self.data_termo_aceite, SDF)
            return _(f"{project} nº {number} - equipada em {date}")
        elif self.data_retorno_assinatura is None:
            date = (
                date_format(self.data_adesao, SDF) if self.data_adesao else ""
            )
            return _(f"{project}, nº {number}, início em {date}")
        if (self.data_retorno_assinatura is not None) and not (
            self.equipada and self.data_termo_aceite is not None
        ):
            date = date_format(self.data_retorno_assinatura, SDF)
            status = self.get_status()
            return _(
                f"{project}, nº {number}, inicio em {date}. Status: {status}"
            )
        if (self.data_retorno_assinatura is not None) and (
            self.equipada and self.data_termo_aceite is not None
        ):
            date = date_format(self.data_retorno_assinatura, SDF)
            equipped_date = date_format(self.data_termo_aceite, SDF)
            return _(
                f"{project}, nº {number}, início em {date} e equipada em "
                f"{equipped_date}. Status: {self.get_status()}"
            )


class EquipamentoPrevisto(models.Model):
    convenio = models.ForeignKey(
        Convenio, on_delete=models.CASCADE, verbose_name=_("convênio")
    )
    equipamento = models.ForeignKey(
        "inventario.Equipamento", on_delete=models.CASCADE
    )
    quantidade = models.PositiveSmallIntegerField(default=1)

    class Meta:
        verbose_name = _("equipamento previsto")
        verbose_name_plural = _("equipamentos previstos")

    def __str__(self):
        return _(f"{self.quantidade} {self.equipamento}(s)")


class Anexo(models.Model):
    convenio = models.ForeignKey(
        Convenio, on_delete=models.CASCADE, verbose_name=_("convênio")
    )
    # caminho no sistema para o documento anexo
    arquivo = models.FileField(
        upload_to="apps/convenios/anexo/arquivo", max_length=500
    )
    descricao = models.CharField(_("descrição"), max_length=70)
    data_pub = models.DateTimeField(
        _("data da publicação do anexo"), default=timezone.localtime
    )

    class Meta:
        ordering = ("-data_pub",)

    def __str__(self):
        return _(f"{self.descricao} publicado em {self.data_pub}")


class UnidadeAdministrativa(models.Model):
    sigla = models.CharField(max_length=10)
    nome = models.CharField(max_length=100)

    def __str__(self):
        return self.sigla


class Tramitacao(models.Model):
    convenio = models.ForeignKey(
        Convenio, on_delete=models.CASCADE, verbose_name=_("convênio")
    )
    unid_admin = models.ForeignKey(
        UnidadeAdministrativa,
        on_delete=models.PROTECT,
        verbose_name=_("Unidade Administrativa"),
    )
    data = models.DateField()
    observacao = models.CharField(
        _("observação"),
        max_length=512,
        null=True,
        blank=True,
    )

    class Meta:
        verbose_name_plural = _("Tramitações")

    def __str__(self):
        in_date = _(f"em {self.data}")  # for focused translation
        result = f"{self.unid_admin} {in_date}"
        if self.observacao:
            result = f"{result} ({self.observacao})"
        return result


class Gescon(models.Model):
    url_gescon = models.URLField(
        _("Webservice Gescon"),
        default=(
            "https://adm.senado.gov.br/gestao-contratos/api/contratos"
            "/busca?especie={s}"
        ),
        help_text=_(
            "Informe o ponto de consulta do webservice do Gescon, "
            "inclusive com a querystring. No ponto onde deve ser "
            "inserida a sigla da subespecie do contrato, use a "
            "marcação {s}.<br/><strong>Por exemplo:</strong> "
            "https://adm.senado.gov.br/gestao-contratos/api/contratos"
            "/busca?especie=<strong>{s}</strong>"
        ),
    )
    subespecies = models.TextField(
        _("Subespécies"),
        default="AC=ACT\nPI=PI\nCN=PML\nTA=PML",
        help_text=_(
            "Informe as siglas das subespécies de contratos que "
            "devem ser pesquisados no Gescon com a sigla "
            "correspondente do projeto no SIGI. Coloque um par de "
            "siglas por linha, no formato SIGLA_GESTON=SIGLA_SIGI. "
            "As siglas não encontradas serão ignoradas."
        ),
    )
    palavras = models.TextField(
        _("Palavras de filtro"),
        default="ILB\nINTERLEGIS",
        help_text=_(
            "Palavras que devem aparecer no campo OBJETO dos dados do "
            "Gescon para identificar se o contrato pertence ao ILB. "
            "<ul><li>Informe uma palavra por linha.</li>"
            "<li>Ocorrendo qualquer uma das palavras, o contrato será "
            "importado.</li></ul>"
        ),
    )
    palavras_excluir = models.TextField(
        _("palavras de exclusão"),
        default="DTCOM",
        help_text=_(
            "Palavras que não podem aparecer no campo OBJETO dos dados do "
            "Gescon."
            "<ul><li>Informe uma palavra por linha.</li>"
            "<li>Ocorrendo qualquer uma das palavras, o contrato será "
            "ignorado.</li></ul>"
        ),
    )
    orgaos_gestores = models.TextField(
        _("Órgãos gestores"),
        default="SCCO",
        help_text=_(
            "Siglas de órgãos gestores que devem aparecer no campo"
            "ORGAOSGESTORESTITULARES"
            "<ul><li>Informe um sigla por linha.</li>"
            "<li>Ocorrendo qualquer uma das siglas, o contrato será "
            "importado.</li></ul>"
        ),
    )
    email = models.EmailField(
        _("E-mail"),
        help_text=_(
            "Caixa de e-mail para onde o relatório diário de "
            "importação será enviado."
        ),
    )
    ultima_importacao = models.TextField(
        _("Resultado da última importação"), blank=True
    )
    checksums = models.JSONField(null=True)

    class Meta:
        verbose_name = _("Configuração do Gescon")
        verbose_name_plural = _("Configurações do Gescon")

    def __str__(self):
        return self.url_gescon

    def save(self, *args, **kwargs):
        self.pk = 1  # Highlander (singleton pattern)
        return super(Gescon, self).save(*args, **kwargs)

    def delete(self, *args, **kwargs):
        pass  # Highlander is immortal

    def add_message(self, msg, save=False):
        self.ultima_importacao += msg + "\n\n"
        if save:
            self.save()
            self.email_report()

    def email_report(self):
        if self.email:
            send_mail(
                subject=_("Relatório de importação GESCON"),
                message=self.ultima_importacao,
                recipient_list=self.email,
                fail_silently=True,
            )
        else:
            self.ultima_importacao += _(
                "\n\n*Não foi definida uma caixa de e-mail nas configurações "
                "do Gescon*"
            )
            self.save()

    def importa_contratos(self):
        """Importa dados de convênios do GESCON

        Returns:
            boolean: Indica se há o que reportar ao usuário (erro ou dados
                     importados/atualizados)
        """

        def mathnames(nome, orgaos, all=False):
            for o, nome_canonico in orgaos:
                ratio = SequenceMatcher(
                    None, to_ascii(nome).lower(), nome_canonico
                ).ratio()
                if ratio > 0.9:
                    yield (o, ratio)

        def get_semelhantes(nome, orgaos, all=False):
            return sorted(
                mathnames(nome, orgaos, all),
                key=lambda m: m[1],
            )

        self.ultima_importacao = ""
        if self.checksums is None:
            self.checksums = {}
        self.add_message(
            _(
                f"Importação iniciada em {timezone.localtime():%d/%m/%Y %H:%M:%S}\n"
                "==========================================================\n"
            )
        )

        if self.palavras == "" or self.orgaos_gestores == "":
            self.add_message(
                _(
                    "Nenhuma palavra de pesquisa ou orgãos "
                    "gestores definidos - processo abortado."
                ),
                True,
            )
            return True

        if self.subespecies == "":
            self.add_message(
                _("Nenhuma subespécie definida - processo abortado."), True
            )
            return True

        if "{s}" not in self.url_gescon:
            self.add_message(
                _(
                    "Falta a marcação {s} na URL para indicar o local onde "
                    "inserir a sigla da subespécia na consulta ao webservice "
                    "- processo abortado."
                ),
                True,
            )
            return True

        palavras = self.palavras.splitlines()
        excludentes = self.palavras_excluir.splitlines()
        orgaos = self.orgaos_gestores.split()
        subespecies = {tuple(s.split("=")) for s in self.subespecies.split()}
        todos_orgaos = [
            (o, f"{to_ascii(o.nome)} - {o.uf_sigla}".lower())
            for o in Orgao.objects.all()
            .order_by()
            .annotate(uf_sigla=F("municipio__uf__sigla"))
        ]

        requests.packages.urllib3.disable_warnings()
        report_user = False
        dominio = get_current_site(None).domain

        for sigla_gescon, sigla_sigi in subespecies:
            self.add_message(_(f"\n**Importando subespécie {sigla_gescon}**"))
            url = self.url_gescon.format(s=sigla_gescon)

            projeto = Projeto.objects.get(sigla=sigla_sigi)

            try:
                response = requests.get(url, verify=False)
            except Exception as e:
                self.add_message(_(f"\tErro ao acessar {url}: {e.message}"))
                report_user = True
                continue

            if not response.ok:
                self.add_message(
                    _(f"\tErro ao acessar {url}: {response.reason}")
                )
                report_user = True
                continue

            if "application/json" not in response.headers.get("Content-Type"):
                self.add_message(
                    _(
                        f"\tResultado da consulta à {url} não "
                        "retornou dados em formato json"
                    )
                )
                report_user = True
                continue

            md5sum = md5(response.text.encode(response.encoding)).hexdigest()
            if (
                sigla_gescon in self.checksums
                and self.checksums[sigla_gescon] == md5sum
            ):
                self.add_message(
                    f"\tDados da subespécie {sigla_gescon} inalterados no "
                    "Gescon. Processamento desnecessário."
                )
                continue

            contratos = response.json()
            Convenio.objects.filter(projeto=projeto).update(erro_gescon=False)

            # Pegar só os contratos que possuem alguma das palavras-chave

            nossos = [
                c
                for c in contratos
                if (
                    any(palavra in c["objeto"] for palavra in palavras)
                    or any(
                        orgao in c["orgaosGestoresTitulares"]
                        for orgao in orgaos
                        if c["orgaosGestoresTitulares"] is not None
                    )
                )
                and not any(palavra in c["objeto"] for palavra in excludentes)
            ]

            self.add_message(
                _(f"\t{len(nossos)} contratos encontrados no Gescon")
            )

            novos = 0
            erros = 0
            alertas = 0
            atualizados = 0

            for contrato in nossos:
                numero = re.sub(
                    r"(\d{4})(\d{4})", r"\1/\2", contrato["numero"].zfill(8)
                )
                sigad = re.sub(
                    r"(\d{5})(\d{6})(\d{4})(\d{2})",
                    r"\1.\2/\3-\4",
                    contrato["processo"].zfill(17),
                )

                if contrato["cnpjCpfFornecedor"]:
                    cnpj = contrato["cnpjCpfFornecedor"].zfill(14)
                    cnpj_masked = mask_cnpj(cnpj)
                else:
                    cnpj = None
                    cnpj_masked = None

                if contrato["nomeFornecedor"]:
                    nome = to_ascii(
                        contrato["nomeFornecedor"]
                        .replace("VEREADORES DE", "")
                        .replace("DO ESTADO", "")
                        .split("-")[0]
                        .split("/")[0]
                        .strip()
                        .replace("  ", " ")
                    )
                else:
                    nome = None

                # Buscar o Convenio pelo NUP #
                convenios = Convenio.objects.filter(
                    projeto=projeto, num_processo_sf=sigad
                )
                if convenios.count() == 0:
                    # Encontrou 0: Pode ser que só exista com o código Gescon
                    convenios = Convenio.objects.filter(
                        Q(projeto=projeto)
                        & Q(Q(num_convenio=numero) | Q(num_processo_sf=numero))
                    )
                if convenios.count() > 1:
                    # Encontrou N: Marcamos todos como erro e reportamos
                    urls = ", ".join(
                        [
                            '<a href="{dominio}{uri}">{id}</a>'.format(
                                dominio=dominio,
                                uri=reverse(
                                    "admin:convenios_convenio_change",
                                    args=[c.id],
                                ),
                                id=c.id,
                            )
                            for c in convenios
                        ]
                    )
                    convenios.update(
                        erro_gescon=True,
                        observacao_gescon=_(
                            "Este convênio possui o mesmo número dos "
                            f"convenios {urls}"
                        ),
                    )
                    self.add_message(
                        _(
                            f"\t* O contrato {numero} no Gescon pode "
                            "ser relacionado aos seguintes convênios "
                            f"do SIGI: {urls}"
                        )
                    )
                    erros += 1
                    # Porém, talvez seja possível ser desambiguado pelo CNPJ do
                    # fornecedor
                    if cnpj_masked is not None:
                        convenios = convenios.filter(
                            casa_legislativa__cnpj=cnpj_masked
                        )
                    if convenios.count() != 1:
                        # Continua ambíguo. Não dá pra fazer nada.
                        continue
                if convenios.count() == 1:
                    # Achou exatamente o único que deveria existir. Basta
                    # atualizar os dados
                    convenio = convenios.get()
                    convenio.projeto = projeto
                    convenio.num_processo_sf = sigad
                    convenio.num_convenio = numero
                    convenio.data_sigad = contrato["assinatura"]
                    convenio.observacao = contrato["objeto"]
                    convenio.data_retorno_assinatura = contrato[
                        "inicioVigencia"
                    ]
                    convenio.data_termino_vigencia = contrato[
                        "terminoVigencia"
                    ]
                    convenio.data_pub_diario = contrato["publicacao"]
                    convenio.atualizacao_gescon = timezone.localtime()
                    convenio.erro_gescon = False
                    convenio.observacao_gescon = ""
                    convenio.id_contrato_gescon = (
                        contrato["codTextoContrato"] or ""
                    )
                    convenio.save()
                    atualizados += 1
                    # Corrigir o CNPJ do órgão se estiver diferente do Gescon
                    # O gescon é um pouquinho mais confiável, por enquanto.
                    if (
                        cnpj_masked
                        and convenio.casa_legislativa.cnpj != cnpj_masked
                    ):
                        convenio.casa_legislativa.cnpj = cnpj_masked
                        convenio.casa_legislativa.save()
                    continue

                # Se chegou aqui, é porque não encontrou o convênio.
                # Um novo convênio precisa ser criado.
                # Primeiro, é preciso identificar qual órgão consta no
                # contrato do Gescon
                if (cnpj is None) and (nome is None):
                    self.add_message(
                        _(
                            f"\t* O contrato {numero} no Gescon não informa "
                            "nem o CNPJ nem o nome do órgão, então não é "
                            "possível importar para o SIGI."
                        )
                    )
                    erros += 1
                    continue
                # Vamos tentar primeiro com o CNPJ
                if cnpj is not None:
                    try:
                        orgao = Orgao.objects.get(cnpj=cnpj_masked)
                    except Orgao.MultipleObjectsReturned:
                        # Pode acontecer de uma câmara usar o mesmo CNPJ
                        # da prefeitura, e ambos terem convênio com o ILB.
                        # Podemos tentar desambiguar pelo nome mais
                        # semelhante.
                        orgao = get_semelhantes(
                            to_ascii(contrato["nomeFornecedor"]).lower(),
                            [
                                (
                                    o,
                                    f"{to_ascii(o.nome)} - {o.uf_sigla}".lower(),
                                )
                                for o in Orgao.objects.filter(cnpj=cnpj_masked)
                                .order_by()
                                .annotate(uf_sigla=F("municipio__uf__sigla"))
                            ],
                            all=True,
                        )[0][0]
                    except Orgao.DoesNotExist:
                        # Encontrou 0: Vamos seguir sem órgao e tentar
                        # encontrar pelo nome logo abaixo
                        orgao = None
                if orgao is None:
                    # Não achou pelo CNPJ. Bora ver se acha por similaridade
                    # do nome
                    if nome is None:
                        # Também não tem nome... então temos que reportar erro
                        self.add_message(
                            _(
                                f"\t* O contrato {numero} no Gescon "
                                f"com NUP sigad {sigad}, fornecedor "
                                f"{cnpj_masked} não pode ser imortado porque "
                                "não é possível identificar o órgão no SIGI. "
                                "Cadastre um órgão com o CNPJ desse "
                                "fornecedor, que na próxima importação este "
                                "contrato será importado."
                            )
                        )
                        erros += 1
                        continue
                    # Tentar primeiro com o nome igual veio do GESCON
                    semelhantes = get_semelhantes(
                        to_ascii(contrato["nomeFornecedor"]).lower(),
                        todos_orgaos,
                    )
                    if not semelhantes:
                        # Não achou, então vamos tentar com o nome limpado
                        semelhantes = get_semelhantes(
                            to_ascii(nome).lower(),
                            todos_orgaos,
                        )
                    if len(semelhantes) > 0:
                        # Encontrou algo semelhante.... bora usar.
                        orgao = semelhantes[0][0]
                    else:
                        # Não encontrou nada parecido. Bora reportar como erro
                        self.add_message(
                            _(
                                f"\t* O contrato {numero} no Gescon "
                                f"com NUP Sigad {sigad}, indica o "
                                f"fornecedor com CNPJ {cnpj_masked} "
                                f"e com o nome {contrato['nomeFornecedor']}, "
                                "que não tem correspondência no SIGI. "
                                "Este convênio precisa ser cadastrado "
                                "manualmente no SIGI para este erro "
                                "parar de acontecer."
                            )
                        )
                        erros += 1
                        continue
                # Não encontrou o órgão... bora reportar o erro
                if orgao is None:
                    # Em teoria, nunca vai cair aqui... mas...
                    self.add_message(
                        _(
                            f"\t* Órgão não encontrado no SIGI ou mais de um "
                            f"órgão encontrado com o mesmo CNPJ ou nome. Favor"
                            f" regularizar o cadastro: "
                            f"CNPJ: {contrato['cnpjCpfFornecedor']}, "
                            f"Nome: {contrato['nomeFornecedor']}"
                        )
                    )
                    erros += 1
                    continue
                else:
                    # Bora criar o convênio
                    convenio = Convenio(
                        casa_legislativa=orgao,
                        projeto=projeto,
                        num_processo_sf=sigad,
                        num_convenio=numero,
                        data_sigi=timezone.localdate(),
                        data_sigad=contrato["assinatura"],
                        observacao=contrato["objeto"],
                        data_retorno_assinatura=contrato["inicioVigencia"],
                        data_termino_vigencia=contrato["terminoVigencia"],
                        data_pub_diario=contrato["publicacao"],
                        atualizacao_gescon=timezone.localtime(),
                        observacao_gescon=_(
                            "Importado integralmente do Gescon"
                        ),
                        id_contrato_gescon=(
                            contrato["codTextoContrato"] or ""
                        ),
                    )
                    convenio.save()
                    novos += 1
                    # Corrigir o CNPJ do órgão se estiver diferente do Gescon
                    # O gescon é um pouquinho mais confiável, por enquanto.
                    if (
                        cnpj_masked
                        and convenio.casa_legislativa.cnpj != cnpj_masked
                    ):
                        convenio.casa_legislativa.cnpj = cnpj_masked
                        convenio.casa_legislativa.save()
                    continue

            if novos or erros or alertas or atualizados:
                report_user = True

            self.add_message(
                _(
                    f"\n\n\t{novos} novos convenios adicionados ao SIGI, "
                    f"{atualizados} atualizados, sendo {alertas} com alertas, e "
                    f"{erros} reportados com erro."
                )
            )
            self.checksums[sigla_gescon] = md5sum

        self.save()
        return report_user

    @classmethod
    def load(cls):
        obj, created = cls.objects.get_or_create(pk=1)
        return obj
