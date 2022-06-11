import calendar
import csv
import locale
from datetime import datetime
from functools import reduce
from typing import OrderedDict
from django.contrib import messages
from django.contrib.admin.sites import site
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
from django.shortcuts import redirect, render, get_object_or_404
from django.template import Template, Context
from django.template.exceptions import TemplateSyntaxError
from django.utils import timezone
from django.utils.text import slugify
from django.utils.translation import (
    to_locale,
    get_language,
    ngettext,
    gettext as _,
)
from django.urls import reverse
from django_weasyprint.utils import django_url_fetcher
from django_weasyprint.views import WeasyTemplateResponse
from docx import Document
from weasyprint import HTML
from sigi.apps.casas.models import Funcionario, Orgao, Presidente
from sigi.apps.convenios.models import Projeto
from sigi.apps.eventos.models import Evento, Equipe, Convite, Modulo, Anexo
from sigi.apps.eventos.forms import (
    SelecionaModeloForm,
    ConviteForm,
    CasaForm,
    FuncionarioForm,
)
from sigi.apps.servidores.models import Servidor


@login_required
def calendario(request):
    mes_pesquisa = int(request.GET.get("mes", timezone.localdate().month))
    ano_pesquisa = int(request.GET.get("ano", timezone.localdate().year))
    formato = request.GET.get("fmt", "cal")
    pdf = bool(request.GET.get("pdf", 0))

    meses = {}
    lang = to_locale(get_language()) + ".UTF-8"
    locale.setlocale(locale.LC_ALL, lang)

    for ano, mes in (
        Evento.objects.exclude(data_inicio=None)
        .values_list("data_inicio__year", "data_inicio__month")
        .order_by("data_inicio__year", "data_inicio__month")
        .distinct("data_inicio__year", "data_inicio__month")
    ):
        if ano in meses:
            meses[ano][mes] = calendar.month_name[mes]
        else:
            meses[ano] = {mes: calendar.month_name[mes]}

    eventos = Evento.objects.filter(
        data_inicio__year=ano_pesquisa, data_inicio__month=mes_pesquisa
    )

    context = site.each_context(request) or {}

    if formato == "cal" or pdf:
        semanas = calendar.Calendar().monthdatescalendar(
            ano_pesquisa, mes_pesquisa
        )
        for semana in semanas:
            for dia in semana:
                if dia.month == mes_pesquisa:
                    semana[dia.weekday()] = (
                        dia.day,
                        [
                            e
                            for e in eventos
                            if e.data_inicio.day
                            <= dia.day
                            <= e.data_termino.day
                        ],
                    )
                else:
                    semana[dia.weekday()] = ("", [])
        context["semanas"] = semanas

    context.update(
        {
            "ano_pesquisa": ano_pesquisa,
            "mes_pesquisa": mes_pesquisa,
            "formato": formato,
            "meses": meses,
            "day_names": calendar.day_abbr,
            "eventos": eventos,
        }
    )

    if pdf:
        context["title"] = _("Calendário de eventos")
        context["pdf"] = True
        return WeasyTemplateResponse(
            filename="calendario_mensal.pdf",
            request=request,
            template="eventos/calendario_pdf.html",
            context=context,
            content_type="application/pdf",
        )
    else:
        return render(request, "eventos/calendario.html", context)


@login_required
def declaracao(request, id):
    if request.method == "POST":
        form = SelecionaModeloForm(request.POST)
        if form.is_valid():
            evento = get_object_or_404(Evento, id=id)
            modelo = form.cleaned_data["modelo"]
            membro = (
                evento.equipe_set.filter(assina_oficio=True).first()
                or evento.equipe_set.first()
            )
            if membro:
                servidor = membro.membro
            else:
                servidor = None
            template_string = (
                """
                {% extends "eventos/declaracao_pdf.html" %}
                {% block text_body %}"""
                + modelo.texto
                + """
                {% endblock %}
                """
            )
            context = Context(
                {
                    "pagesize": modelo.formato,
                    "pagemargin": modelo.margem,
                    "evento": evento,
                    "servidor": servidor,
                    "data": timezone.localdate(),
                }
            )
            string = Template(template_string).render(context)
            # return HttpResponse(string)
            response = HttpResponse(
                headers={
                    "Content-Type": "application/pdf",
                    "Content-Disposition": 'attachment; filename="declaração.pdf"',
                }
            )
            pdf = HTML(
                string=string,
                url_fetcher=django_url_fetcher,
                encoding="utf-8",
                base_url=request.build_absolute_uri("/"),
            )
            pdf.write_pdf(target=response)
            return response
    else:
        form = SelecionaModeloForm()

    context = site.each_context(request)
    context["form"] = form
    context["evento_id"] = id

    return render(request, "eventos/seleciona_modelo.html", context)


def evento(request, id):
    context = site.each_context(request)
    evento = get_object_or_404(Evento, id=id)
    anexo_id = request.GET.getlist("anexo_id", None)
    if anexo_id:
        context["anexos"] = evento.anexo_set.filter(id__in=anexo_id)
        context["active"] = "anexos"
    else:
        context["anexos"] = evento.anexo_set.all()

    context["evento"] = evento
    context["fields"] = [
        "tipo_evento",
        "descricao",
        "virtual",
        "publico_alvo",
        "data_inicio",
        "data_termino",
        "carga_horaria",
        "casa_anfitria",
        "municipio",
        "local",
    ]
    context["convite_fields"] = [
        "casa",
        "servidor",
        "data_convite",
        "aceite",
        "participou",
        "nomes_participantes",
    ]
    context["anexo_fields"] = ["descricao", "data_pub", "arquivo"]

    return render(request, "eventos/evento.html", context)


def convida_casa(request, evento_id, casa_id):
    def processa_paragrafos(paragrafos, context):
        for paragrafo in paragrafos:
            run_final = None
            for run in paragrafo.runs:
                if run_final is None:
                    run_final = run
                else:
                    run_final.text += run.text
                    run.text = ""
                try:
                    run_final.text = Template(run_final.text).render(context)
                    run_final = None
                except TemplateSyntaxError:
                    pass

    if not request.user.servidor:
        messages.error(
            request, _("Você não é servidor, não pode registrar convites")
        )
        return redirect(reverse("eventos-evento", args=[evento_id]))

    evento = get_object_or_404(Evento, id=evento_id)
    casa = get_object_or_404(Orgao, id=casa_id)

    projetos = Projeto.objects.exclude(modelo_minuta="")

    if evento.convite_set.filter(casa=casa).exists():
        convite = evento.convite_set.get(casa=casa)
    else:
        convite = Convite(
            evento=evento,
            casa=casa,
            servidor=request.user.servidor,
            data_convite=timezone.localdate(),
        )

    presidente = casa.presidente or Funcionario(
        casa_legislativa=casa, setor="presidente"
    )
    contato = casa.contato_interlegis or Funcionario(
        casa_legislativa=casa, setor="contato_interlegis"
    )

    if request.method == "POST":
        form_convite = ConviteForm(request.POST, instance=convite)
        form_casa = CasaForm(request.POST, request.FILES, instance=casa)
        form_presidente = FuncionarioForm(
            request.POST, instance=presidente, prefix="presidente"
        )
        form_contato = FuncionarioForm(
            request.POST, instance=contato, prefix="contato"
        )

        if all(
            [
                form_convite.is_valid(),
                form_casa.is_valid(),
                form_presidente.is_valid(),
                form_contato.is_valid(),
            ]
        ):
            contato = form_contato.save()
            presidente = form_presidente.save()
            casa = form_casa.save()
            convite = form_convite.save()

            proj_id = request.POST.get("save", "")

            if proj_id:
                convite.anexo_set.all().delete()
                query_str = ""
                projeto = get_object_or_404(Projeto, id=proj_id)
                if projeto.texto_oficio:
                    oficio = gerar_anexo(
                        casa,
                        presidente,
                        contato,
                        path=request.build_absolute_uri("/"),
                        nome=f"Ofício de solicitação de {projeto.sigla}",
                        modelo="oficio_padrao.html",
                        texto=projeto.texto_oficio,
                    )
                    oficio.evento = evento
                    oficio.save()
                    query_str += f"anexo_id={oficio.id}&"
                if projeto.modelo_minuta:
                    doc = Document(projeto.modelo_minuta.path)
                    if casa.tipo.sigla == "CM":
                        ente = (
                            f"Município de {casa.municipio.nome}, "
                            f"{casa.municipio.uf.sigla}"
                        )
                    else:
                        ente = f"Estado de {casa.municipio.uf.nome}"
                    doc_context = Context(
                        {
                            "evento": evento,
                            "casa": casa,
                            "presidente": presidente,
                            "contato": contato,
                            "data": timezone.localdate(),
                            "ente": ente,
                            "doravante": casa.tipo.nome.split(" ")[0],
                        }
                    )
                    processa_paragrafos(doc.paragraphs, doc_context)
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                processa_paragrafos(
                                    cell.paragraphs,
                                    doc_context,
                                )
                    nome = f"Minuta de {projeto.sigla} da {casa.nome}"[:70]
                    minuta = Anexo(descricao=nome, evento=evento)
                    minuta.arquivo.name = slugify(nome) + ".docx"
                    doc.save(minuta.arquivo.path)
                    minuta.save()
                    query_str += f"anexo_id={minuta.id}"

                return redirect(evento.get_absolute_url() + "?" + query_str)
            else:
                return redirect(evento.get_absolute_url())
        else:
            messages.error(_("Preencha corretamente o convite"))
    else:
        form_convite = ConviteForm(instance=convite)
        form_casa = CasaForm(instance=casa)
        form_presidente = FuncionarioForm(
            instance=presidente, prefix="presidente"
        )
        form_contato = FuncionarioForm(instance=contato, prefix="contato")

    context = site.each_context(request)
    context.update(
        {
            "form_convite": form_convite,
            "form_casa": form_casa,
            "form_presidente": form_presidente,
            "form_contato": form_contato,
            "evento": evento,
            "convite": convite,
            "casa": casa,
            "presidente": presidente,
            "contato": contato,
            "projetos": projetos,
        }
    )

    return render(request, "eventos/convida_casa.html", context)


def gerar_anexo(casa, presidente, contato, path, modelo, nome, texto):
    template_string = (
        f'{{% extends "eventos/{modelo}" %}}'
        "{% load pdf %}"
        f"{{% block text_body %}}{texto}{{% endblock %}}"
    )
    context = Context(
        {
            "evento": evento,
            "casa": casa,
            "presidente": presidente,
            "contato": contato,
            "data": timezone.localdate(),
            "doravante": casa.tipo.nome.split(" ")[0],
        }
    )
    string = Template(template_string).render(context)
    pdf = HTML(
        string=string,
        url_fetcher=django_url_fetcher,
        encoding="utf-8",
        base_url=path,
    )
    nome = (nome + f" da {casa.nome}")[:70]
    anexo = Anexo(descricao=nome)
    anexo.arquivo.name = slugify(nome) + ".pdf"
    f = anexo.arquivo.open("wb")
    pdf.write_pdf(target=f)
    f.flush()
    f.close()
    return anexo


@login_required
def alocacao_equipe(request):
    ano_pesquisa = int(request.GET.get("ano", timezone.localdate().year))
    mes_pesquisa = int(request.GET.get("mes", 0))
    semana_pesquisa = int(request.GET.get("semana", 0))
    formato = request.GET.get("fmt", "html")

    lang = to_locale(get_language()) + ".UTF-8"
    locale.setlocale(locale.LC_ALL, lang)

    dados = []
    eventos = Evento.objects.exclude(status="C").prefetch_related("equipe_set")

    num_cols = 12

    if mes_pesquisa > 0:
        semanas = [
            [s[0], s[-1]]
            for s in calendar.Calendar().monthdatescalendar(
                ano_pesquisa, mes_pesquisa
            )
        ]
        num_cols = len(semanas)
        if semana_pesquisa > 0:
            dias = calendar.Calendar().monthdatescalendar(
                ano_pesquisa, mes_pesquisa
            )[semana_pesquisa - 1]
            num_cols = len(dias)
            eventos = eventos.filter(
                data_inicio__gte=dias[0], data_inicio__lte=dias[-1]
            )
        else:
            eventos = eventos.filter(
                data_inicio__gte=semanas[0][0], data_inicio__lte=semanas[-1][-1]
            )
    else:
        eventos = eventos.filter(data_inicio__year=ano_pesquisa)

    for evento in eventos:
        for p in evento.equipe_set.all():
            registro = None
            for r in dados:
                if r[0] == p.membro.pk:
                    registro = r
                    break
            if not registro:
                if semana_pesquisa > 0:
                    registro = [
                        p.membro.pk,
                        p.membro.nome_completo,
                        OrderedDict([(dia, []) for dia in dias]),
                    ]
                else:
                    registro = [
                        p.membro.pk,
                        p.membro.nome_completo,
                        [{"dias": 0, "eventos": 0} for __ in range(num_cols)],
                    ]
                dados.append(registro)

            if mes_pesquisa > 0:
                if semana_pesquisa > 0:
                    for dia in dias:
                        if (
                            evento.data_inicio.date()
                            <= dia
                            <= evento.data_termino.date()
                        ):
                            registro[2][dia].append(evento)
                else:
                    for idx, [inicio, fim] in enumerate(semanas):
                        if inicio <= evento.data_inicio.date() <= fim:
                            registro[2][idx]["dias"] += (
                                min(fim, evento.data_termino.date())
                                - evento.data_inicio.date()
                            ).days + 1
                            registro[2][idx]["eventos"] += 1
                        elif inicio <= evento.data_termino.date() <= fim:
                            registro[2][idx]["dias"] += (
                                min(fim, evento.data_termino.date())
                                - evento.data_inicio.date()
                            ).days + 1
                            registro[2][idx]["eventos"] += 1
            else:
                registro[2][evento.data_inicio.month - 1]["dias"] += (
                    evento.data_termino - evento.data_inicio
                ).days + 1
                registro[2][evento.data_inicio.month - 1]["eventos"] += 1

    dados.sort(key=lambda x: x[1])

    meses = list(calendar.month_abbr)[1:]
    linhas = []

    if semana_pesquisa:
        linhas = [
            [registro[1]] + list(registro[2].values()) for registro in dados
        ]
    else:
        for r in dados:
            r[2].append(
                reduce(
                    lambda x, y: {
                        "dias": x["dias"] + y["dias"],
                        "eventos": x["eventos"] + y["eventos"],
                    },
                    r[2],
                )
            )
            linhas.append(
                [r[1]]
                + [
                    _(
                        ngettext("%(dias)s dia", "%(dias)s dias", d["dias"])
                        + " em "
                        + ngettext(
                            "%(eventos)s evento",
                            "%(eventos)s eventos",
                            d["eventos"],
                        )
                    )
                    % d
                    if d["dias"] > 0 or d["eventos"] > 0
                    else ""
                    for d in r[2]
                ]
            )

    context = site.each_context(request) or {}
    context.update(
        {
            "anos": Evento.objects.exclude(data_inicio=None)
            .order_by("data_inicio__year")
            .distinct("data_inicio__year")
            .values_list("data_inicio__year", flat=True),
            "ano_pesquisa": ano_pesquisa,
            "linhas": linhas,
        }
    )
    if mes_pesquisa > 0:
        context["mes_pesquisa"] = mes_pesquisa
        context["meses"] = meses
        if semana_pesquisa > 0:
            cabecalho = [_("Servidor")] + dias
            context["semana_pesquisa"] = semana_pesquisa
            context["eventos"] = eventos
        else:
            cabecalho = (
                [_("Servidor")]
                + [
                    _(f"de {inicio:%d/%m} a {fim:%d/%m}")
                    for inicio, fim in semanas
                ]
                + ["total"]
            )
    else:
        cabecalho = [_("Servidor")] + meses + ["total"]

    context["cabecalho"] = cabecalho

    if formato == "pdf":
        context["title"] = _("Alocação de equipe")
        context["pdf"] = True
        return WeasyTemplateResponse(
            # filename="alocacao_equipe.pdf",
            request=request,
            template="eventos/alocacao_equipe_pdf.html",
            context=context,
            content_type="application/pdf",
        )
    elif formato == "csv":
        response = HttpResponse(content_type="text/csv")
        response[
            "Content-Disposition"
        ] = 'attachment; filename="alocacao_equipe_%s.csv"' % (ano_pesquisa,)
        writer = csv.writer(response)
        writer.writerow(cabecalho)
        writer.writerows(linhas)
        return response

    return render(request, "eventos/alocacao_equipe.html", context)


# # Views e functions para carrinho de exportação

# def query_ordena(qs, o):
#     from sigi.apps.eventos.admin import EventoAdmin
#     list_display = EventoAdmin.list_display
#     order_fields = []

#     for order_number in o.split('.'):
#         order_number = int(order_number)
#         order = ''
#         if order_number != abs(order_number):
#             order_number = abs(order_number)
#             order = '-'
#         order_fields.append(order + list_display[order_number - 1])
#     qs = qs.order_by(*order_fields)
#     return qs

# def get_for_qs(get, qs):
#     kwargs = {}
#     for k, v in get.iteritems():
#         if str(k) not in ('page', 'pop', 'q', '_popup', 'o', 'ot'):
#             kwargs[str(k)] = v
#     qs = qs.filter(**kwargs)
#     if 'o' in get:
#         qs = query_ordena(qs, get['o'])
#     return qs

# def carrinhoOrGet_for_qs(request):
#     if 'carrinho_eventos' in request.session:
#         ids = request.session['carrinho_eventos']
#         qs = Evento.objects.filter(pk__in=ids)
#     else:
#         qs = Evento.objects.all()
#         if request.GET:
#             qs = get_for_qs(request.GET, qs)
#     return qs

# def adicionar_eventos_carrinho(request, queryset=None, id=None):
#     if request.method == 'POST':
#         ids_selecionados = request.POST.getlist('_selected_action')
#         if 'carrinho_eventos' not in request.session:
#             request.session['carrinho_eventos'] = ids_selecionados
#         else:
#             lista = request.session['carrinho_eventos']
#             # Verifica se id já não está adicionado
#             for id in ids_selecionados:
#                 if id not in lista:
#                     lista.append(id)
#             request.session['carrinho_eventos'] = lista

# @login_required
# def visualizar_carrinho(request):
#     qs = carrinhoOrGet_for_qs(request)
#     paginator = Paginator(qs, 100)

#     try:
#         page = int(request.GET.get('page', '1'))
#     except ValueError:
#         page = 1

#     try:
#         paginas = paginator.page(page)
#     except (EmptyPage, InvalidPage):
#         paginas = paginator.page(paginator.num_pages)

#     carrinhoIsEmpty = not('carrinho_eventos' in request.session)

#     return render(
#         request,
#         'eventos/carrinho.html',
#         {
#             'carIsEmpty': carrinhoIsEmpty,
#             'paginas': paginas,
#             'query_str': '?' + request.META['QUERY_STRING']
#         }
#     )

# @login_required
# def excluir_carrinho(request):
#     if 'carrinho_eventos' in request.session:
#         del request.session['carrinho_eventos']
#         messages.info(request, 'O carrinho foi esvaziado')
#     return HttpResponseRedirect('../../')

# @login_required
# def deleta_itens_carrinho(request):
#     if request.method == 'POST':
#         ids_selecionados = request.POST.getlist('_selected_action')
#         removed = 0
#         if 'carrinho_eventos' in request.session:
#             lista = request.session['carrinho_eventos']
#             for item in ids_selecionados:
#                 lista.remove(item)
#                 removed += 1
#             if lista:
#                 request.session['carrinho_eventos'] = lista
#             else:
#                 del lista
#                 del request.session['carrinho_eventos']
#         messages.info(request, "{0} itens removidos do carrinho".format(removed))
#     return HttpResponseRedirect('.')

# @login_required
# def export_csv(request):
#     def rm_rows(lista,reg):
#         for a in lista:
#             if a in lista:
#                 reg.pop(a,None)
#             else:
#                 pass

#     def serialize(r, field):
#         value = (getattr(r, 'get_{0}_display'.format(field.name), None) or
#                  getattr(r, field.name, ""))
#         if callable(value):
#             value = value()
#         if value is None:
#             value = ""
#         return unicode(value).encode('utf8')

#     eventos = carrinhoOrGet_for_qs(request)
#     eventos.select_related('equipe', 'convite')

#     if not eventos:
#         messages.info(request, _("Nenhum evento a exportar"))
#         return HttpResponseRedirect('../')

#     max_equipe = max([e.equipe_set.count() for e in eventos])

#     mun_casa = 'Município da Casa Anfitriã'.encode('utf8')
#     uf_casa = 'UF da Casa Anfitriã'.encode('utf8')
#     reg_casa = 'Região da Casa Anfitriã'.encode('utf8')

#     head = [f.verbose_name.encode('utf8') for f in Evento._meta.fields]
#     head.extend([mun_casa, uf_casa, reg_casa])
#     head.extend([f.verbose_name.encode('utf8')+"_{0}".format(i+1)
#         for i in range(max_equipe) for f in Equipe._meta.fields
#         if f.name not in ('id', 'evento')])
#     head.extend([f.verbose_name.encode('utf8') for f in Convite._meta.fields
#         if f.name not in ('id', 'evento')])
#     head.extend([f.verbose_name.encode('utf8') for f in Modulo._meta.fields
#         if f.name not in ('id', 'evento')])

#     response = HttpResponse(content_type='text/csv')
#     response['Content-Disposition'] = 'attachment; filename=eventos.csv'
#     rm_list = ['Descrição do evento', 'Local do evento', 'Público alvo', 'Motivo do cancelamento', 'Descrição do módulo']

#     for a in head:
#         if 'Observações_' in a:
#             rm_list.append(a)

#     for a in rm_list:
#         if a in head:
#             head.remove(a)
#         else:
#             pass
#     writer = csv.DictWriter(response, fieldnames=head)
#     writer.writeheader()

#     for evento in eventos:
#         reg = {f.verbose_name.encode('utf8'): serialize(evento, f)
#                for f in Evento._meta.fields}
#         if evento.casa_anfitria is None:
#             reg[mun_casa] = ""
#             reg[uf_casa] = ""
#             reg[reg_casa] = ""
#         else:
#             reg[mun_casa] = evento.casa_anfitria.municipio.nome.encode('utf8')
#             reg[uf_casa] = evento.casa_anfitria.municipio.uf.sigla.\
#                 encode('utf8')
#             reg[reg_casa] = evento.casa_anfitria.municipio.uf.\
#                 get_regiao_display().encode('utf8')

#         idx = 1
#         for membro in evento.equipe_set.all():
#             reg.update(
#                 {
#                     "{0}_{1}".format(f.verbose_name.encode('utf8'), idx):
#                         serialize(membro, f) for f in Equipe._meta.fields
#                         if f.name not in ('id', 'evento')
#                 }
#             )
#             idx += 1
#         for convite in evento.convite_set.all():
#             reg.update(
#                 {f.verbose_name.encode('utf8'): serialize(convite, f)
#                  for f in Convite._meta.fields
#                  if f.name not in ('id', 'evento')}
#             )
#             rm_rows(rm_list,reg)
#             writer.writerow(reg)

#         if evento.convite_set.count() == 0:
#             rm_rows(rm_list,reg)

#             writer.writerow(reg)

#     return response
